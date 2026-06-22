"""Deterministic DOCX generator (五期 B+ 多类型扩展).

The model supplies a *structured document* (title + sections, each with
paragraphs and/or bullets); this module renders it into a real ``.docx`` using a
fixed, code-controlled style (``python-docx``). The model never writes rendering
code (A 铁律).

Design goals mirror ``pptx_generator``: deterministic, dependency-light
(``python-docx`` + stdlib), defensive (normalize loose / oversized output), and
generic / open-source safe (single env-overridable accent).

Public API:
    normalize_doc(raw) -> dict          # validate + clamp into a canonical shape
    build_docx(doc) -> bytes            # render canonical doc → .docx bytes
    safe_filename(title, ext="docx")    # re-exported from file_gen_common
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

import file_gen_common as common

# --- Limits (clamp model output) ----------------------------------------------
MAX_SECTIONS = 60
MAX_PARAS_PER_SECTION = 60
MAX_BULLETS_PER_SECTION = 60
MAX_TITLE_CHARS = 200
MAX_SUBTITLE_CHARS = 300
MAX_HEADING_CHARS = 200
MAX_PARA_CHARS = 5000
MAX_BULLET_CHARS = 1000


def _str_list(raw: Any, limit_chars: int, max_items: int) -> list[str]:
    if isinstance(raw, str):
        raw = [raw]
    if not isinstance(raw, list):
        return []
    out: list[str] = []
    for item in raw:
        text = common.clean_text(item, limit_chars)
        if text:
            out.append(text)
        if len(out) >= max_items:
            break
    return out


def normalize_doc(raw: Any) -> dict[str, Any]:
    """Validate + clamp loose model output into a canonical document.

    Canonical shape::

        {
          "title": str,
          "subtitle": str,                 # may be ""
          "sections": [
            {"heading": str, "paragraphs": [str], "bullets": [str]},
            ...
          ],
        }

    Never raises — returns something renderable.
    """
    if not isinstance(raw, dict):
        raw = {}

    title = common.clean_text(raw.get("title") or raw.get("name"), MAX_TITLE_CHARS) or "未命名文档"
    subtitle = common.clean_text(raw.get("subtitle") or raw.get("author") or "", MAX_SUBTITLE_CHARS)

    raw_sections = raw.get("sections") or raw.get("blocks")
    if not isinstance(raw_sections, list):
        raw_sections = []

    sections: list[dict[str, Any]] = []
    for s in raw_sections:
        if isinstance(s, str):
            # bare string → a paragraph-only section with no heading
            if s.strip():
                sections.append({"heading": "", "paragraphs": [common.clean_text(s, MAX_PARA_CHARS)], "bullets": []})
        elif isinstance(s, dict):
            heading = common.clean_text(s.get("heading") or s.get("title"), MAX_HEADING_CHARS)
            paragraphs = _str_list(
                s.get("paragraphs") or s.get("body") or s.get("text"),
                MAX_PARA_CHARS, MAX_PARAS_PER_SECTION,
            )
            bullets = _str_list(
                s.get("bullets") or s.get("points") or s.get("items"),
                MAX_BULLET_CHARS, MAX_BULLETS_PER_SECTION,
            )
            if not heading and not paragraphs and not bullets:
                continue
            sections.append({"heading": heading, "paragraphs": paragraphs, "bullets": bullets})
        if len(sections) >= MAX_SECTIONS:
            break

    # Also accept a flat ``paragraphs`` / ``bullets`` payload (no sections).
    if not sections:
        flat_paras = _str_list(raw.get("paragraphs") or raw.get("body"), MAX_PARA_CHARS, MAX_PARAS_PER_SECTION)
        flat_bullets = _str_list(raw.get("bullets"), MAX_BULLET_CHARS, MAX_BULLETS_PER_SECTION)
        if flat_paras or flat_bullets:
            sections = [{"heading": "", "paragraphs": flat_paras, "bullets": flat_bullets}]

    if not sections:
        sections = [{"heading": "", "paragraphs": [], "bullets": []}]

    return {"title": title, "subtitle": subtitle, "sections": sections}


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------


def build_docx(doc_data: Any) -> bytes:
    """Render a (canonical or loose) document into ``.docx`` bytes."""
    doc_data = normalize_doc(doc_data)
    accent = RGBColor.from_string(common.accent_hex())
    muted = RGBColor.from_string("8A8A8A")

    document = Document()

    # Title
    h = document.add_heading(level=0)
    run = h.add_run(doc_data["title"])
    run.font.color.rgb = accent

    # Subtitle (muted)
    if doc_data.get("subtitle"):
        p = document.add_paragraph()
        r = p.add_run(doc_data["subtitle"])
        r.italic = True
        r.font.size = Pt(12)
        r.font.color.rgb = muted

    for section in doc_data["sections"]:
        if section["heading"]:
            sh = document.add_heading(level=1)
            sr = sh.add_run(section["heading"])
            sr.font.color.rgb = accent
        for para in section["paragraphs"]:
            document.add_paragraph(para)
        for bullet in section["bullets"]:
            # 'List Bullet' is a built-in python-docx style (always present).
            document.add_paragraph(bullet, style="List Bullet")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def safe_filename(title: str, ext: str = "docx") -> str:
    return common.safe_filename(title, ext, fallback="document")


if __name__ == "__main__":  # pragma: no cover — local smoke test
    import sys

    sample = {
        "title": "Q4 营收复盘报告",
        "subtitle": "经营分析组 · 2026-06",
        "sections": [
            {
                "heading": "整体表现",
                "paragraphs": ["Q4 总营收同比 +18.4%,创历史新高,主要由线上渠道与新店爬坡共同拉动。"],
                "bullets": ["线上渠道贡献 32%", "门店数净增 47 家", "客单价同比 +5.2%"],
            },
            {
                "heading": "下一步计划",
                "bullets": ["加密华南布点", "华北门店运营诊断", "线上会员复购专项"],
            },
        ],
    }
    data = build_docx(sample)
    out = sys.argv[1] if len(sys.argv) > 1 else "/tmp/sample.docx"
    with open(out, "wb") as f:
        f.write(data)
    reopened = Document(io.BytesIO(data))
    print(f"OK: {len(data)} bytes, paragraphs={len(reopened.paragraphs)} → {out}")
    print(f"filename: {safe_filename(sample['title'])}")
